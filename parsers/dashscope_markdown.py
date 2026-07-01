"""
DashScope (Bailian) markdown parser — sends documents to qwen-plus
for structured markdown extraction with heading hierarchy preserved.

This serves as a fallback when Marker (local GPU) is unavailable,
and as the primary parser for .docx files when the user prefers
Bailian over Marker.
"""

import logging
from typing import Optional

from openai import OpenAI

from config import (
    DASHSCOPE_API_KEY,
    DASHSCOPE_BASE_URL,
    DASHSCOPE_MARKDOWN_MODEL,
)
from models.schemas import ParsedDocument, PageBlock
from parsers.base import BaseParser

logger = logging.getLogger(__name__)

MARKDOWN_EXTRACT_PROMPT = """你是一个专业的文档解析助手。请将以下文档内容转换为结构化Markdown格式。

要求：
1. 严格保留原始文档的标题层级（#、##、###）
2. 保留所有表格，使用 | ... | ... | 格式
3. 保留所有列表、编号、缩进
4. 在每个页面边界标注 [PAGE: 页码]
5. 保留页码信息以便后续引用
6. 直接输出Markdown，不要添加任何额外解释

文档内容如下：
---
{document_text}"""


class DashScopeMarkdownParser(BaseParser):
    """Extract structured markdown via Bailian qwen-plus for digital documents."""

    def __init__(self, api_key: Optional[str] = None, model: Optional[str] = None):
        self._api_key = api_key or DASHSCOPE_API_KEY
        self._model = model or DASHSCOPE_MARKDOWN_MODEL
        self._client: Optional[OpenAI] = None

    @property
    def name(self) -> str:
        return "dashscope_markdown"

    def _ensure_client(self):
        if self._client is None:
            if not self._api_key:
                raise ValueError(
                    "DASHSCOPE_API_KEY not set. "
                    "Cannot use Bailian for markdown extraction."
                )
            self._client = OpenAI(
                api_key=self._api_key,
                base_url=DASHSCOPE_BASE_URL,
            )

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        self._ensure_client()

        # For .docx, extract text with python-docx first
        # For .pdf, extract text with pypdf first
        # Then send to qwen-plus for structuring
        raw_text = self._extract_raw_text(file_path, filename)

        logger.info(
            f"[{self.name}] Sending {len(raw_text)} chars to {self._model} "
            f"for markdown extraction"
        )

        response = self._client.chat.completions.create(
            model=self._model,
            messages=[
                {
                    "role": "user",
                    "content": MARKDOWN_EXTRACT_PROMPT.format(
                        document_text=raw_text[:120000]  # Truncate to context limit
                    ),
                }
            ],
            max_tokens=8000,
            temperature=0.1,
        )

        markdown = response.choices[0].message.content or ""

        if not markdown.strip():
            raise ValueError(
                f"Bailian markdown extraction returned empty result for '{filename}'."
            )

        # Extract page blocks from [PAGE: N] markers
        pages = self._split_by_page_markers(markdown)

        return ParsedDocument(
            markdown=markdown,
            pages=pages,
            toc=[],
            parser_used=self.name,
        )

    def _extract_raw_text(self, file_path: str, filename: str) -> str:
        """Extract raw text as input for Bailian structuring."""
        import os
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    parts.append(f"[PAGE: {i+1}]\n{text}")
            return "\n\n".join(parts)

        elif ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Detect heading-like styles
                    if para.style and para.style.name and "Heading" in para.style.name:
                        level = para.style.name.replace("Heading", "").strip()
                        try:
                            lv = int(level)
                            parts.append(f"{'#' * lv} {para.text}")
                        except ValueError:
                            parts.append(para.text)
                    else:
                        parts.append(para.text)
            return "\n\n".join(parts)

        else:
            # Plain text — try common encodings
            encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
            for enc in encodings:
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        return f.read()
                except (UnicodeDecodeError, UnicodeError):
                    continue
            raise ValueError(f"Cannot read '{filename}' with any supported encoding.")

    @staticmethod
    def _split_by_page_markers(markdown: str) -> list[PageBlock]:
        """Split markdown into PageBlocks using [PAGE: N] markers."""
        import re
        # Split on [PAGE: N] markers
        pattern = r"\[PAGE:\s*(\d+)\]"
        parts = re.split(pattern, markdown)

        pages = []
        # parts[0] = text before first marker (if any)
        # parts[1] = page num, parts[2] = content, parts[3] = page num, ...
        i = 0
        if parts and not re.match(pattern, parts[0]):
            # Text before any page marker → page 1
            if parts[0].strip():
                pages.append(PageBlock(page_num=1, markdown=parts[0].strip()))
            i = 1

        while i + 1 < len(parts):
            page_num = int(parts[i])
            content = parts[i + 1].strip()
            if content:
                pages.append(PageBlock(page_num=page_num, markdown=content))
            i += 2

        return pages
